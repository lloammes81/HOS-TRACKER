from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

def heading(text, size=13, bold=True, center=False, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p

def body(text, size=11, bold=False, italic=False, indent=0, space_after=6,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold   = bold
    run.italic = italic
    return p

def hr():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '4B5563')
    pBdr.append(bottom)
    pPr.append(pBdr)

# Encabezado tribunal
heading("IN THE COUNTY COURT OF THE FIFTEENTH JUDICIAL CIRCUIT", size=11, center=True, space_before=0, space_after=2)
heading("IN AND FOR PALM BEACH COUNTY, FLORIDA", size=11, center=True, space_before=0, space_after=2)
heading("CRIMINAL DIVISION", size=11, center=True, space_before=0, space_after=10)
hr()

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after  = Pt(2)
run = p.add_run("STATE OF FLORIDA,")
run.bold = True
run.font.size = Pt(11)

body("                        Plaintiff,", size=11)
p2 = doc.add_paragraph()
p2.paragraph_format.space_after = Pt(2)
run2 = p2.add_run("vs.")
run2.font.size = Pt(11)

body("LLOAMMES FABRE,", size=11, bold=True)
body("                        Defendant (Pro Se).", size=11)
body("", size=6)

body("Numeros de Caso:", size=10, bold=True)
cases = [
    ("50-2026-MO-000750-AXXX-MB", "6013347"),
    ("50-2026-MO-000751-AXXX-MB", "6013380"),
    ("50-2026-MO-000752-AXXX-MB", "6013326"),
    ("50-2026-MO-000753-AXXX-MB", "6013345"),
    ("50-2026-MO-000754-AXXX-MB", "6013379"),
]
for cn, cit in cases:
    body("  Caso: " + cn + "   |   Citacion: " + cit, size=10, indent=0.3)

hr()
heading("MOCION DE DESESTIMACION Y ESCRITO DE DEFENSA", size=14, center=True, space_before=14, space_after=4)
heading("(PRO SE  --  AUDIENCIA PRE-TRIAL  --  11 DE MAYO DE 2026)", size=11, center=True, space_before=0, space_after=10)

# I
heading("I. INTRODUCCION Y COMPARECENCIA", size=12, space_before=14)
body(
    "El suscrito, LLOAMMES FABRE, comparece ante este Honorable Tribunal en calidad de "
    "Demandado Pro Se, actuando en defensa propia de conformidad con el derecho que le asiste "
    "bajo la Constitucion del Estado de Florida y la Constitucion de los Estados Unidos de America. "
    "El Demandado respetuosamente solicita que se desestimen las presentes citaciones de transito, "
    "por cuanto las mismas fueron expedidas con errores materiales y/o procesales que las hacen "
    "invalidas e inaplicables como cuestion de derecho.", size=11)

# II
heading("II. DECLARACION DE HECHOS", size=12, space_before=14)
body(
    "1.  El dia 31 de marzo de 2026, el Demandado, LLOAMMES FABRE, recibio cinco (5) citaciones "
    "de transito emitidas en el Condado de Palm Beach, Florida.", size=11)
body(
    "2.  El Demandado afirma que dichas citaciones contienen errores en su emision, incluyendo "
    "pero no limitado a: inexactitudes en los datos registrados, falta de fundamento observable "
    "para la infraccion imputada, y/o deficiencias en el procedimiento seguido por el agente emisor.", size=11)
body(
    "3.  Ante la existencia de dichos errores, el Demandado presento su impugnacion de las "
    "citaciones oportunamente, siendo convocado a la presente audiencia pre-trial del 11 de mayo "
    "de 2026 a las 11:00 a.m. via Zoom (Meeting ID: 964 4865 4396).", size=11)
body(
    "4.  El Demandado no ha renunciado a ningun derecho procesal o constitucional y comparece "
    "expresamente reservando todos sus derechos.", size=11)

# III
heading("III. ARGUMENTOS LEGALES", size=12, space_before=14)

heading("A. Error Material en la Expedicion de las Citaciones", size=11, bold=True, space_before=10, space_after=4)
body(
    "Conforme al Seccion 318.14, Florida Statutes, una citacion de transito debe ser expedida "
    "de manera precisa y completa para ser valida. Cuando una citacion contiene errores "
    "materiales -- tales como datos incorrectos sobre el lugar, el estatuto supuestamente "
    "infringido, la descripcion del vehiculo o las circunstancias del incidente -- dicho "
    "documento pierde su fuerza legal. El Demandado afirma que las cinco (5) citaciones "
    "impugnadas presentan tales deficiencias, lo que justifica su desestimacion.", size=11)

heading("B. Derecho al Discovery -- Regla 3.220 Fla. R. Crim. P.", size=11, bold=True, space_before=10, space_after=4)
body(
    "De conformidad con la Regla 3.220 de las Florida Rules of Criminal Procedure, el Demandado "
    "tiene derecho a solicitar y obtener todos los materiales relevantes en posesion del Estado, "
    "incluyendo sin limitacion:", size=11)

items = [
    "Las notas de campo y el reporte completo del agente emisor.",
    "Cualquier grabacion de camara corporal (bodycam) o camara del vehiculo policial (dashcam) relacionada con el incidente del 31 de marzo de 2026.",
    "Los registros de calibracion y certificacion de cualquier dispositivo de medicion utilizado.",
    "El historial de certificaciones y entrenamiento del agente emisor.",
    "Cualquier otro documento o evidencia que el Estado pretenda utilizar en contra del Demandado.",
]
for item in items:
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(item)
    run.font.size = Pt(11)

body(
    "El Demandado solicita formalmente que dicho material le sea entregado con anterioridad "
    "a cualquier juicio, de conformidad con la ley aplicable.", size=11, space_after=8)

heading("C. Insuficiencia de Evidencia / Ausencia de Causa Probable", size=11, bold=True, space_before=10, space_after=4)
body(
    "El Estado tiene la carga de probar cada elemento de la infraccion mas alla de toda duda "
    "razonable. El Demandado afirma que el agente emisor no contaba con base legal suficiente "
    "para expedir las citaciones, y que la evidencia en poder del Estado es insuficiente para "
    "sostener los cargos. Sin la produccion de evidencia objetiva, confiable y debidamente "
    "certificada, el Estado no puede cumplir con su carga probatoria.", size=11)

heading("D. Debido Proceso -- Decimocuarta Enmienda / Art. I Sec. 9 Const. FL", size=11, bold=True, space_before=10, space_after=4)
body(
    "La Decimocuarta Enmienda de la Constitucion de los Estados Unidos y el Articulo I, Seccion 9 "
    "de la Constitucion del Estado de Florida garantizan al Demandado el derecho al debido proceso "
    "legal. La expedicion de citaciones con errores materiales y sin base factual comprobable viola "
    "este derecho fundamental. Permitir que tales citaciones prosperen sin revision rigurosa "
    "constituiria una denegacion del debido proceso.", size=11)

heading("E. Consolidacion de los Cinco (5) Casos", size=11, bold=True, space_before=10, space_after=4)
body(
    "Los cinco (5) casos aqui relacionados surgieron del mismo evento ocurrido el 31 de marzo de "
    "2026, involucrando al mismo Demandado, al mismo agente y las mismas circunstancias de hecho. "
    "Por tanto, cualquier defecto legal o procesal que invalide una de las citaciones se extiende "
    "a todas ellas. El Demandado solicita que este Tribunal las trate de manera consolidada y "
    "resuelva sobre todas ellas en forma conjunta.", size=11)

# IV
heading("IV. SOLICITUDES AL TRIBUNAL", size=12, space_before=14)
body("Por las razones expuestas, el Demandado solicita respetuosamente a este Honorable Tribunal:", size=11)

solicitudes = [
    "DESESTIMAR las cinco (5) citaciones en su totalidad, por defectos materiales y/o insuficiencia de evidencia.",
    "En la alternativa, ORDENAR al Estado que produzca de inmediato todo el material de discovery solicitado en la Seccion III-B.",
    "FIJAR una fecha de juicio en caso de que el Estado mantenga los cargos, preservando todos los derechos del Demandado.",
    "Cualquier otro remedio que este Tribunal considere justo y apropiado.",
]
for i, s in enumerate(solicitudes, 1):
    body(str(i) + ".  " + s, size=11, indent=0.3)

# V - Guia
heading("V. GUIA PRACTICA PARA LA AUDIENCIA (USO PERSONAL)", size=12, space_before=14)
body(
    "Frases sugeridas para la audiencia de hoy. Hable con calma y "
    "dirija siempre su palabra al juez como 'Su Senoria'.",
    size=11, italic=True)

frases = [
    ("Al abrir su turno:",
     "Su Senoria, buenos dias. Me presento Pro Se en representacion propia. Deseo informar al Tribunal que contesto todas las citaciones por cuanto fueron expedidas con errores en su emision y carecen de base legal suficiente."),
    ("Si el fiscal ofrece un acuerdo:",
     "Agradezco la oferta, pero en este momento no estoy en posicion de aceptarla. Solicito que se me proporcione el discovery completo antes de tomar cualquier decision."),
    ("Para solicitar discovery:",
     "Su Senoria, solicito formalmente el discovery, incluyendo las notas del agente, video de dashcam o bodycam, y los registros de calibracion de cualquier dispositivo utilizado, de conformidad con la Regla 3.220 de las Florida Rules of Criminal Procedure."),
    ("Si desea ir a juicio:",
     "Su Senoria, deseo ejercer mi derecho a un juicio y preservo todos mis derechos constitucionales y procesales."),
]
for titulo, frase in frases:
    body(titulo, size=11, bold=True, space_after=2)
    body(frase, size=11, italic=True, indent=0.3, space_after=10)

# Firma
hr()
body("", size=8)
body("Fecha: 11 de mayo de 2026", size=11)
body("", size=8)
body("Respetuosamente sometido,", size=11)
body("", size=8)
body("________________________________________", size=11)
body("LLOAMMES FABRE", size=11, bold=True)
body("2892 Tennis Club Dr, Apt 602", size=11)
body("West Palm Beach, FL 33417", size=11)
body("Demandado Pro Se", size=11)

heading("CERTIFICADO DE SERVICIO", size=11, space_before=18)
body(
    "Certifico que una copia del presente escrito fue entregada al Tribunal el dia "
    "11 de mayo de 2026 mediante comparecencia en la audiencia virtual.", size=11)
body("", size=8)
body("________________________________________", size=11)
body("LLOAMMES FABRE", size=11, bold=True)

output = "/home/user/HOS-TRACKER/Defensa_Citaciones_Fabre_2026.docx"
doc.save(output)
print("OK: " + output)
